"""Regenerate and replace fig_3_8_rag in BAO_CAO docx (find table by section heading)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.shared import Cm
from docx.table import Table
from docx.text.paragraph import Paragraph

from generate_architecture_figures import OUT_DIR, fig_3_8_rag

DOCX = Path(__file__).parent / "BAO_CAO_DO_AN_TOT_NGHIEP_v3.docx"
WIDTH = Cm(15)
SECTION_MARK = "3.7. Kiến trúc RAG Pipeline"


def iter_blocks(doc):
    for child in doc.element.body:
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def replace_table_image(table: Table, image_path: Path) -> None:
    cell = table.rows[0].cells[0]
    cell.text = ""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=WIDTH)


def find_rag_table(doc: Document) -> tuple[int, Table] | None:
    seen_section = False
    for idx, block in enumerate(iter_blocks(doc)):
        if isinstance(block, Paragraph):
            if block.text.strip() == SECTION_MARK:
                seen_section = True
        elif isinstance(block, Table) and seen_section:
            return idx, block
    return None


def main() -> None:
    fig_3_8_rag()
    print(f"Generated {OUT_DIR / 'fig_3_8_rag.png'}")
    doc = Document(str(DOCX))
    found = find_rag_table(doc)
    if not found:
        raise SystemExit(f"RAG table after '{SECTION_MARK}' not found")
    idx, table = found
    replace_table_image(table, OUT_DIR / "fig_3_8_rag.png")
    try:
        doc.save(str(DOCX))
        out = DOCX
    except PermissionError:
        alt = DOCX.with_name("BAO_CAO_DO_AN_TOT_NGHIEP_v3_rag_fix.docx")
        doc.save(str(alt))
        out = alt
        print(f"Note: {DOCX.name} locked — saved {alt.name}; close Word then copy over v3")
    print(f"Updated Hình 3.8 (block {idx}) → {out.name}")


if __name__ == "__main__":
    main()
