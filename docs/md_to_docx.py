"""Convert BAO_CAO markdown to formatted Word (.docx).

- Times New Roman 12pt body, Heading 1/2/3 giữ cỡ lớn hơn
- Hình/sơ đồ: trang riêng + khung viền (không bị cắt khi qua trang)
- Bỏ qua mục MỤC LỤC (sinh viên tự tạo trong Word)
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

MD_PATH = Path(__file__).parent / "BAO_CAO_DO_AN_TOT_NGHIEP_v2.md"
OUT_PATH = Path(__file__).parent / "BAO_CAO_DO_AN_TOT_NGHIEP_v3.docx"

FIGURE_TITLE_RE = re.compile(r"^\*\*(Hình\s+[\d.]+[^*]*)\*\*", re.I)
SKIP_LINE = (
    re.compile(r"^<div\b", re.I),
    re.compile(r"^</div>", re.I),
    re.compile(r"^<br", re.I),
)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)

    sizes = {1: 16, 2: 14, 3: 13}
    for level, size in sizes.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for edge, val in kwargs.items():
        tag = f"w:{edge}"
        element = OxmlElement(tag)
        for key in ("sz", "val", "color", "space"):
            if key in val:
                element.set(qn(f"w:{key}"), str(val[key]))
        tc_pr.append(element)


def bordered_box(doc: Document, content: str, *, mono: bool = True) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "000000"},
        bottom={"sz": 12, "val": "single", "color": "000000"},
        start={"sz": 12, "val": "single", "color": "000000"},
        end={"sz": 12, "val": "single", "color": "000000"},
    )
    cell.width = Cm(15)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(content)
    run.font.name = "Consolas" if mono else "Times New Roman"
    run.font.size = Pt(9 if mono else 12)
    if not mono:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.add_paragraph()


def add_figure_page(doc: Document, title: str, body: str) -> None:
    """Trang riêng: khung hình trước, chú thích Hình X.X phía dưới (chuẩn đồ án VN)."""
    doc.add_page_break()
    bordered_box(doc, body, mono=True)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(12)
    cr = cp.add_run(title)
    cr.bold = True
    cr.font.name = "Times New Roman"
    cr.font.size = Pt(12)
    cr._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text.strip()


def add_rich_paragraph(doc: Document, text: str, style: str | None = None, size: int = 12) -> None:
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*.+?\*\*|\*.+?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [strip_md_inline(c.strip()) for c in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-+:?", c.replace(" ", "")) for c in row):
            rows.append(row)
        i += 1
    return rows, i


def add_md_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            c = table.rows[r_idx].cells[c_idx]
            c.text = cell
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.add_paragraph()


def convert() -> None:
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()
    setup_styles(doc)

    skip_muc_luc = False
    in_code = False
    is_figure = False
    figure_title = ""
    code_buf: list[str] = []
    pending_figure_title: str | None = None
    i = 0

    while i < len(lines):
        line = lines[i]
        raw = line.strip()

        if raw == "## MỤC LỤC":
            skip_muc_luc = True
            i += 1
            continue
        if skip_muc_luc:
            if raw.startswith("## ") and raw != "## MỤC LỤC":
                skip_muc_luc = False
            else:
                i += 1
                continue

        if any(p.search(raw) for p in SKIP_LINE):
            i += 1
            continue

        m = FIGURE_TITLE_RE.match(raw)
        if m and not in_code:
            pending_figure_title = strip_md_inline(m.group(1))
            i += 1
            continue

        if raw.startswith("```"):
            if in_code:
                block = "\n".join(code_buf)
                if is_figure and figure_title:
                    add_figure_page(doc, figure_title, block)
                else:
                    bordered_box(doc, block, mono=True)
                code_buf = []
                in_code = False
                is_figure = False
                figure_title = ""
                pending_figure_title = None
            else:
                in_code = True
                if pending_figure_title:
                    is_figure = True
                    figure_title = pending_figure_title
                    pending_figure_title = None
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if pending_figure_title and raw:
            pending_figure_title = None

        if not raw:
            i += 1
            continue

        if raw == "---":
            doc.add_page_break()
            i += 1
            continue

        if raw.startswith("|"):
            rows, i = parse_table(lines, i)
            add_md_table(doc, rows)
            continue

        if raw.startswith("# ") and not raw.startswith("## "):
            doc.add_heading(strip_md_inline(raw[2:]), level=1)
            i += 1
            continue
        if raw.startswith("## ") and not raw.startswith("### "):
            doc.add_heading(strip_md_inline(raw[3:]), level=2)
            i += 1
            continue
        if raw.startswith("### "):
            doc.add_heading(strip_md_inline(raw[4:]), level=3)
            i += 1
            continue

        if raw.startswith("- ") or raw.startswith("* "):
            add_rich_paragraph(doc, strip_md_inline(raw[2:]), style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s", raw):
            add_rich_paragraph(doc, strip_md_inline(re.sub(r"^\d+\.\s", "", raw)), style="List Number")
            i += 1
            continue

        add_rich_paragraph(doc, raw)
        i += 1

    doc.save(OUT_PATH)
    print(f"Created: {OUT_PATH}")


if __name__ == "__main__":
    convert()
