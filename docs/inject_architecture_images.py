"""Replace ASCII architecture diagrams in BAO_CAO docx with PNG images."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.shared import Cm
from docx.table import Table
from docx.text.paragraph import Paragraph

from generate_architecture_figures import BLOCK_IMAGE_MAP, OUT_DIR, generate_all

DOCX = Path(__file__).parent / "BAO_CAO_DO_AN_TOT_NGHIEP_v3.docx"
IMAGE_WIDTH = Cm(15)


def iter_block_items(parent):
    for child in parent.element.body:
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def replace_table_with_image(table: Table, image_path: Path) -> None:
    cell = table.rows[0].cells[0]
    cell.text = ""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = p.paragraph_format.space_after = 0
    run = p.add_run()
    run.add_picture(str(image_path), width=IMAGE_WIDTH)


def inject() -> None:
    images = generate_all()
    doc = Document(str(DOCX))
    replaced = 0
    idx = 0
    for block in iter_block_items(doc):
        if isinstance(block, Table) and idx in BLOCK_IMAGE_MAP:
            fname = BLOCK_IMAGE_MAP[idx]
            img_path = images[idx] if idx in images else OUT_DIR / fname
            if img_path.exists():
                replace_table_with_image(block, img_path)
                replaced += 1
                print(f"Injected {fname} at block {idx}")
        idx += 1
    doc.save(str(DOCX))
    print(f"Saved {DOCX} — replaced {replaced} diagram(s)")


if __name__ == "__main__":
    inject()
