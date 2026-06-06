"""Insert localhost UI screenshots into BAO_CAO docx after section headings."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

DOCX = Path(__file__).parent / "BAO_CAO_DO_AN_TOT_NGHIEP_v3.docx"
OUT_DOCX = Path(__file__).parent / "BAO_CAO_DO_AN_TOT_NGHIEP_v4.docx"
UI_DIR = Path(__file__).parent / "images" / "ui"
WIDTH = Cm(15)

# block index -> list of (filename, caption)
INSERT_AFTER: dict[int, list[tuple[str, str]]] = {
    467: [("fig_5_1_landing.png", "Hình 5.1 — Giao diện Landing Page (localhost:3000)")],
    475: [
        ("fig_5_2_login.png", "Hình 5.2a — Giao diện đăng nhập (/login)"),
        ("fig_5_2_register.png", "Hình 5.2b — Giao diện đăng ký (/register)"),
    ],
    483: [("fig_5_3_chat.png", "Hình 5.3 — Giao diện Chat AI (/chat)")],
    496: [("fig_5_4_recommendations.png", "Hình 5.4 — Dashboard gợi ý (/recommendations)")],
    499: [
        ("fig_5_5_trends.png", "Hình 5.5a — Trang xu hướng (/trends)"),
        ("fig_5_5_trends_detail.png", "Hình 5.5b — Chi tiết xu hướng (/trends/[id])"),
    ],
    510: [("fig_5_6_outfits.png", "Hình 5.6 — Outfit đã lưu (/outfits)")],
    513: [("fig_5_7_profile.png", "Hình 5.7 — Hồ sơ người dùng (/profile)")],
    515: [("fig_5_8_admin.png", "Hình 5.8 — Bảng quản trị (/admin)")],
    520: [("fig_6_1_docker.png", "Hình 6.1 — Triển khai Docker Compose (docker compose ps)")],
    527: [
        ("fig_6_2_performance.png", "Hình 6.2a — Dashboard metrics & biểu đồ hiệu năng"),
        ("fig_6_2_curl_latency.png", "Hình 6.2b — Đo latency API endpoint (dev local)"),
    ],
    532: [("fig_6_3_swagger.png", "Hình 6.3 — Swagger UI kiểm thử API (localhost:8000/docs)")],
    661: [("fig_5_1_landing.png", "Hình G.1 — Landing Page")],
    663: [("fig_5_3_chat.png", "Hình G.2 — Chat Interface")],
    665: [
        ("fig_5_5_trends.png", "Hình G.3a — Trends Page"),
        ("fig_5_5_trends_detail.png", "Hình G.3b — Trend Detail Page"),
    ],
}


def iter_blocks(doc):
    for child in doc.element.body:
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def paragraph_has_image(p: Paragraph) -> bool:
    for run in p.runs:
        if run._element.findall(".//" + qn("a:blip")):
            return True
    return False


def next_paragraph_has_image(p: Paragraph) -> bool:
    nxt = p._p.getnext()
    if nxt is not None and nxt.tag.endswith("}p"):
        np = Paragraph(nxt, p._parent)
        return paragraph_has_image(np)
    return False


def insert_after(paragraph: Paragraph) -> Paragraph:
    from docx.oxml import OxmlElement

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def add_figure_after(paragraph: Paragraph, image_path: Path, caption: str) -> None:
    p = insert_after(paragraph)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = p.paragraph_format.space_after = 0
    run = p.add_run()
    run.add_picture(str(image_path), width=WIDTH)

    cp = insert_after(p)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Cm(0.2)
    cp.paragraph_format.space_after = Cm(0.4)
    cr = cp.add_run(caption)
    cr.bold = True
    cr.font.name = "Times New Roman"
    cr.font.size = Pt(12)
    cr._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def inject() -> None:
    doc = Document(str(DOCX))
    targets: dict[int, Paragraph] = {}
    for idx, block in enumerate(iter_blocks(doc)):
        if isinstance(block, Paragraph) and idx in INSERT_AFTER:
            targets[idx] = block

    inserted = 0
    for block_idx in sorted(INSERT_AFTER.keys(), reverse=True):
        para = targets.get(block_idx)
        if para is None:
            print(f"Skip missing block {block_idx}")
            continue
        if next_paragraph_has_image(para):
            print(f"Skip block {block_idx} — image already present")
            continue
        for fname, caption in reversed(INSERT_AFTER[block_idx]):
            img = UI_DIR / fname
            if not img.exists():
                print(f"Missing {img}")
                continue
            add_figure_after(para, img, caption)
            inserted += 1
            print(f"Inserted {fname} after block {block_idx}")

    # Update appendix C note if present
    for p in doc.paragraphs:
        if p.text.strip().startswith("Chèn vào vị trí Hình 5.1"):
            p.text = (
                "Đã chèn screenshot UI thật (localhost) cho Hình 5.1–5.8, Hình 6.1–6.3 "
                "và Phụ lục G. Ảnh gốc: docs/images/ui/."
            )
            break

    out = OUT_DOCX
    try:
        doc.save(str(DOCX))
        out = DOCX
    except PermissionError:
        doc.save(str(OUT_DOCX))
        print(f"Note: {DOCX.name} is locked — saved to {OUT_DOCX.name}")
    print(f"Saved {out.name} — {inserted} figure(s) added")


if __name__ == "__main__":
    inject()
